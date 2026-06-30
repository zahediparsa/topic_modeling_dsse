from pathlib import Path
import json
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path('outputs/week2/mapreduce')
summary = json.loads((ROOT/'summary.json').read_text(encoding='utf-8'))
lda = pd.read_csv(ROOT/'lda_final_topics.csv')
ber = pd.read_csv(ROOT/'bertopic_topic_summary.csv')
coh = pd.read_csv(ROOT/'lda_coherence_scores.csv')
props = pd.read_csv(ROOT/'lda_issue_topics.csv')
topic_cols = [c for c in props if c.startswith('topic_')]
multi = (props[topic_cols] >= .2).sum(axis=1)

NAVY='17365D'; BLUE='2E74B5'; PALE='E8EEF5'; GRAY='F2F4F7'; WHITE='FFFFFF'; MUTED='666666'
doc = Document(); sec=doc.sections[0]
sec.page_width=Inches(8.5); sec.page_height=Inches(11)
sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
sec.header_distance=sec.footer_distance=Inches(.492)

def rgb(h): return RGBColor.from_string(h)
def set_font(run, name='Calibri', size=None, bold=None, color=None, italic=None):
    run.font.name=name; run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'),name); run._element.rPr.rFonts.set(qn('w:hAnsi'),name)
    if size: run.font.size=Pt(size)
    if bold is not None: run.bold=bold
    if color: run.font.color.rgb=rgb(color)
    if italic is not None: run.italic=italic

normal=doc.styles['Normal']; normal.font.name='Calibri'; normal.font.size=Pt(11); normal.font.color.rgb=rgb('222222')
normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.10
for name,size,before,after,color in [('Heading 1',16,16,8,BLUE),('Heading 2',13,12,6,BLUE),('Heading 3',12,8,4,'1F4D78')]:
    st=doc.styles[name]; st.font.name='Calibri'; st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=rgb(color)
    st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True
for name in ('List Bullet','List Number'):
    st=doc.styles[name]; st.font.name='Calibri'; st.font.size=Pt(11); st.paragraph_format.left_indent=Inches(.5); st.paragraph_format.first_line_indent=Inches(-.25); st.paragraph_format.space_after=Pt(8); st.paragraph_format.line_spacing=1.167

header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
set_font(header.add_run('DSSE Assignment 3 | Week 2'), size=9, color=MUTED)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
set_font(footer.add_run('Apache MapReduce - Topic Modeling Analysis'), size=9, color=MUTED)

def set_repeat_header(row):
    trpr=row._tr.get_or_add_trPr(); el=OxmlElement('w:tblHeader'); el.set(qn('w:val'),'true'); trpr.append(el)
def set_cant_split(row): row._tr.get_or_add_trPr().append(OxmlElement('w:cantSplit'))
def shade(cell, fill):
    shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); cell._tc.get_or_add_tcPr().append(shd)
def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tcpr=cell._tc.get_or_add_tcPr(); tcMar=tcpr.first_child_found_in('w:tcMar') or OxmlElement('w:tcMar')
    if tcMar.getparent() is None: tcpr.append(tcMar)
    for edge,val in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=OxmlElement(f'w:{edge}'); node.set(qn('w:w'),str(val)); node.set(qn('w:type'),'dxa'); tcMar.append(node)
def fix_table(table, widths):
    table.autofit=False; table.alignment=WD_TABLE_ALIGNMENT.LEFT
    tblpr=table._tbl.tblPr; tblw=tblpr.first_child_found_in('w:tblW'); tblw.set(qn('w:w'),'9360'); tblw.set(qn('w:type'),'dxa')
    ind=OxmlElement('w:tblInd'); ind.set(qn('w:w'),'120'); ind.set(qn('w:type'),'dxa'); tblpr.append(ind)
    grid=table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths:
        gc=OxmlElement('w:gridCol'); gc.set(qn('w:w'),str(width)); grid.append(gc)
    for ri,row in enumerate(table.rows):
        set_cant_split(row)
        if ri==0: set_repeat_header(row)
        for ci,cell in enumerate(row.cells):
            cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; set_cell_margins(cell)
            tcw=cell._tc.get_or_add_tcPr().first_child_found_in('w:tcW'); tcw.set(qn('w:w'),str(widths[ci])); tcw.set(qn('w:type'),'dxa')
            for p in cell.paragraphs:
                p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.05
                for r in p.runs: set_font(r,size=9.3,color='222222')
        if ri==0:
            for cell in row.cells:
                shade(cell,GRAY)
                for r in cell.paragraphs[0].runs: set_font(r,size=9.3,bold=True,color=NAVY)
def add_table(headers, rows, widths):
    t=doc.add_table(rows=1, cols=len(headers)); t.style='Table Grid'
    for i,h in enumerate(headers): t.rows[0].cells[i].text=str(h)
    for vals in rows:
        cells=t.add_row().cells
        for i,v in enumerate(vals): cells[i].text=str(v)
    fix_table(t,widths); return t
def add_caption(text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(4)
    set_font(p.add_run(text),size=9,italic=True,color=MUTED)

# memo masthead
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(4)
set_font(p.add_run('TECHNICAL REPORT'),size=10,bold=True,color=BLUE)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4)
set_font(p.add_run('Week 2: LDA and BERTopic'),size=25,bold=True,color=NAVY)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(16)
set_font(p.add_run('Topic modeling of 587 Apache MapReduce design issues'),size=14,color=MUTED)
add_table(['Field','Value'], [('Project','Apache MapReduce'),('Corpus','587 issue summaries + descriptions'),('Models','LDA and BERTopic'),('Reproducibility','Random seed 42; issue-level assignments retained')],[1800,7560])

doc.add_heading('Executive result',1)
doc.add_paragraph('The refined LDA model selected 9 topics at the highest tested c_v coherence (0.4095). BERTopic produced 16 focused semantic clusters and treated 193 issues as outliers. Together, the models reveal recurring concerns around execution lifecycle, scheduling, data movement, storage formats, resource limits, reliability, and security.')

doc.add_heading('1. Method',1)
doc.add_heading('1.1 LDA iteration 1 - baseline',2)
doc.add_paragraph('The first iteration used the Week 1 stemmed tokens, 10 topics, and low document-topic and topic-word priors (alpha = 0.01; beta = 0.01). This sparse prior reflects the expectation that one issue discusses few topics and each topic is expressed by a limited term set. The baseline still contained project-specific and generic terms, so several topics overlapped.')
doc.add_heading('1.2 LDA iteration 2 - refined vocabulary',2)
doc.add_paragraph('High-frequency generic/project tokens were removed. Terms with a confident semantic role were replaced by ontology classes: Component, Connector, Data, Solution, and separate quality-attribute classes for performance, reliability, security, scalability, and maintainability. Models with 3 through 10 topics were then fitted under identical priors.')
doc.add_heading('1.3 BERTopic',2)
doc.add_paragraph('BERTopic embedded the original summary and description text with all-MiniLM-L6-v2, reduced the embeddings with UMAP, clustered them with HDBSCAN, and extracted topic terms with class-based TF-IDF. A leaf-based density configuration was used after the initial broad cluster was found too coarse. Outliers were retained as outliers rather than force-assigned.')

doc.add_page_break()
doc.add_heading('2. LDA optimization',1)
doc.add_picture(str(ROOT/'lda_coherence.png'),width=Inches(6.3))
add_caption('Figure 1. c_v coherence for 3-10 refined LDA topics. Nine topics achieved the maximum score.')
rows=[(int(r.n_topics),f'{r.c_v_coherence:.4f}', 'Selected' if int(r.n_topics)==9 else '') for _,r in coh.iterrows()]
add_table(['Topics','c_v coherence','Decision'],rows,[1800,3000,4560])
doc.add_paragraph('The 9-topic model was selected because its c_v coherence of 0.4095 was the highest tested value. The drop at 10 topics suggests that an additional topic fragmented coherent themes rather than adding useful separation.')
doc.add_picture(str(ROOT/'lda_topic_counts.png'),width=Inches(6.3))
add_caption('Figure 2. Issues counted by dominant LDA topic.')

doc.add_heading('3. Final LDA topics',1)
lda_rows=[]
for _,r in lda.iterrows():
    terms=', '.join(str(r.top_terms).split(', ')[:8])
    lda_rows.append((f"L{int(r.topic)}",r.label,int(r.issue_count),terms))
add_table(['ID','Interpretation','Issues','Top keywords'],lda_rows,[700,3000,900,4760])
add_caption('Table 1. Issue counts use each issue\'s dominant LDA topic; keywords are ordered by topic weight.')
doc.add_heading('Topic proportions per issue',2)
doc.add_paragraph(f"At a probability threshold of 0.20, issues contain an average of {multi.mean():.2f} LDA topics. {int((multi==1).sum())} issues have one topic above the threshold and {int((multi>1).sum())} have multiple topics. The mean dominant-topic probability is {props.dominant_probability.mean():.3f}, showing that LDA commonly captures a strong primary concern plus secondary concerns.")

doc.add_heading('4. BERTopic results',1)
doc.add_picture(str(ROOT/'bertopic_topic_counts.png'),width=Inches(6.3))
add_caption('Figure 3. BERTopic cluster sizes, excluding the outlier class.')
doc.add_paragraph('BERTopic found 16 interpretable clusters covering 394 issues. The remaining 193 issues (32.9%) were marked as outliers by HDBSCAN. This is expected in density-based topic modeling and is preferable to assigning isolated or ambiguous issues to an artificial topic.')
ber_non=ber[ber.topic!=-1]
ber_rows=[]
for _,r in ber_non.iterrows():
    terms=', '.join(str(r.top_terms).split(', ')[:7])
    ber_rows.append((f"B{int(r.topic)}",r.label,int(r.issue_count),terms))
add_table(['ID','Interpretation','Issues','Top keywords'],ber_rows,[700,3200,800,4660])
add_caption('Table 2. BERTopic cluster definitions and issue counts; the outlier class is excluded.')

doc.add_heading('5. Answer to Research Question 1',1)
doc.add_paragraph('RQ1: What topics emerge from LDA and BERTopic, what are their common keywords, and how many issues discuss each topic?')
doc.add_paragraph('LDA produced nine broad themes. The largest were Job state, history and submission (125 issues), HDFS/RAID/input streams (106), and Trackers/slots/heartbeats (87). BERTopic separated these broad themes into sixteen narrower clusters. Its largest non-outlier clusters were YARN integration/application clients (49), Sort/merge/reduce pipeline (43), and Schedulers/queues/preemption (39). Complete keywords and counts appear in Tables 1 and 2 and in the accompanying CSV files.')
doc.add_heading('Cross-model interpretation',2)
for text in [
    'Execution management appears in both models: LDA execution/tracker topics correspond to BERTopic clusters for YARN clients, task memory, job history, and heartbeats.',
    'Data processing is broad in LDA but separates in BERTopic into sort/merge, shuffle transfer, streaming, distributed cache, input formats, RAID blocks, split sizing, and serialization.',
    'Scheduling is stable across methods: LDA scheduling/coordination themes align with the BERTopic scheduler, queue, preemption, and workload-trace clusters.',
    'Reliability and security are visible in LDA ontology-aware topics; BERTopic isolates a specific failure/retry cluster but security is distributed mainly across YARN/client and configuration clusters.',
]: doc.add_paragraph(text,style='List Bullet')
doc.add_heading('Usefulness and limitations',2)
doc.add_paragraph('The refined LDA topics are useful for broad retrieval categories and allow multi-topic issue proportions. BERTopic is stronger for fine-grained technical search but leaves ambiguous issues unclustered. Topic labels are analyst interpretations of the reported keywords, not ground-truth classes. Results are stochastic in general; the pipeline fixes seeds and saves models and assignments to make this run reproducible.')

doc.add_heading('6. Reproducibility and outputs',1)
for text in [
    'week2_topic_modeling.py - end-to-end LDA and BERTopic pipeline',
    'lda_iteration1_topics.csv - baseline 10-topic result',
    'lda_coherence_scores.csv and lda_coherence.png - optimization evidence',
    'lda_final_topics.csv and lda_issue_topics.csv - final definitions and issue proportions',
    'bertopic_topic_summary.csv and bertopic_issue_topics.csv - semantic clusters and assignments',
    'lda_final_model.joblib and bertopic_model/ - saved models',
]: doc.add_paragraph(text,style='List Bullet')

doc.add_heading('Sources',1)
sources=[
    'BERTopic algorithm and modular pipeline: https://maartengr.github.io/BERTopic/algorithm/algorithm.html',
    'BERTopic API and output attributes: https://maartengr.github.io/BERTopic/api/bertopic.html',
    'Gensim c_v coherence model: https://radimrehurek.com/gensim/models/coherencemodel.html',
    'Scikit-learn LDA API: https://scikit-learn.org/stable/modules/generated/sklearn.decomposition.LatentDirichletAllocation.html',
]
for s in sources:
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.18); p.paragraph_format.first_line_indent=Inches(-.18); p.paragraph_format.space_after=Pt(3)
    set_font(p.add_run('• ' + s),size=8.5,color=MUTED)

ROOT.mkdir(parents=True,exist_ok=True)
doc.save(ROOT/'Week_2_Report_MapReduce.docx')
