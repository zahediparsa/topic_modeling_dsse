from pathlib import Path
import json
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path('outputs/week1/mapreduce')
summary = json.loads((ROOT / 'summary.json').read_text(encoding='utf-8'))

doc = Document()
section = doc.sections[0]
section.top_margin = section.bottom_margin = Inches(0.75)
section.left_margin = section.right_margin = Inches(0.85)
styles = doc.styles
styles['Normal'].font.name = 'Aptos'
styles['Normal'].font.size = Pt(10.5)
for name, size, color in [('Title', 24, '17365D'), ('Heading 1', 16, '17365D'), ('Heading 2', 12, '2F5597')]:
    styles[name].font.name = 'Aptos Display'
    styles[name].font.size = Pt(size)
    styles[name].font.color.rgb = RGBColor.from_string(color)

title = doc.add_paragraph(style='Title')
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run('Week 1: Issue Data and Vocabulary')
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run('Apache MapReduce - DSSE Assignment 3')
r.bold = True; r.font.color.rgb = RGBColor.from_string('4472C4')

doc.add_heading('1. Scope and completion', level=1)
doc.add_paragraph(
    'The complete list of 587 MapReduce issues from Issues.xlsx was processed. '
    'Issue records were downloaded from the Apache Jira REST API and cached locally. '
    'All assigned issues were retrieved successfully; no downloads failed.'
)
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Shading Accent 1'
table.rows[0].cells[0].text = 'Measure'; table.rows[0].cells[1].text = 'Result'
metrics = [
    ('Assigned/downloaded issues', f"{summary['downloaded_issues']} / {summary['assigned_issues']}"),
    ('Issues with a parent', str(summary['issues_with_parent'])),
    ('Processed tokens', f"{summary['total_tokens']:,}"),
    ('Unique vocabulary terms', f"{summary['vocabulary_size']:,}"),
    ('Document-term matrix', f"{summary['matrix_shape'][0]} x {summary['matrix_shape'][1]}"),
    ('Non-zero matrix entries', f"{summary['matrix_nonzero_entries']:,}"),
]
for label, value in metrics:
    cells = table.add_row().cells; cells[0].text = label; cells[1].text = value

doc.add_heading('2. Jira data collection', level=1)
doc.add_paragraph(
    'For every issue ID, the pipeline requested all Jira fields and retained the complete raw fields '
    'inside the JSONL dataset. A compact CSV contains the most useful analysis fields: issue ID, three '
    'design-decision labels, parent ID, summary, type, status, priority, resolution, dates, comment counts, '
    'and attachment count. The parent field was absent for most issues, as expected.'
)
doc.add_paragraph(
    'The supplied bot-author list is built into the pipeline. The authoritative Jira total comment count '
    'is stored for each issue; human-comment counts are also calculated for the comments included in the issue response.'
)

doc.add_heading('3. Text preprocessing', level=1)
for text in [
    'Concatenated each issue summary and description.',
    'Converted text to lowercase and tokenized alphabetic/alphanumeric software terms.',
    'Removed English stop words, short tokens, numeric-only tokens, and URL-like tokens.',
    'Applied English Snowball stemming.',
    'Preserved the ordered token list for every issue in issues_processed.jsonl.',
]:
    doc.add_paragraph(text, style='List Bullet')

doc.add_heading('4. Vocabulary results', level=1)
top = summary['top_30_tokens'][:20]
vt = doc.add_table(rows=1, cols=4)
vt.style = 'Light Shading Accent 1'
for i, h in enumerate(['Rank', 'Token', 'Frequency', 'Interpretation']): vt.rows[0].cells[i].text = h
project_specific = {'job', 'task', 'reduc', 'map', 'hadoop', 'mapreduc', 'jobtrack', 'tasktrack', 'mapper'}
generic = {'use', 'need', 'current', 'new', 'number', 'implement', 'support'}
for rank, (token, frequency) in enumerate(top, 1):
    interpretation = 'Project-specific candidate' if token in project_specific else ('Generic candidate' if token in generic else 'Keep/review')
    cells = vt.add_row().cells
    for cell, value in zip(cells, [str(rank), token, str(frequency), interpretation]): cell.text = value

doc.add_heading('5. Candidates for refinement', level=1)
doc.add_paragraph(
    'The most frequent terms reveal several likely refinements before LDA. Project names and platform vocabulary '
    '(hadoop, mapreduc) may dominate topics without describing a design concern. Generic action words '
    '(use, need, current, new, implement, support) are also weak discriminators and are candidates for an '
    'expanded domain stop-word list. Terms such as job, task, mapper, reducer, JobTracker, TaskTracker, cluster, '
    'and node should not be deleted automatically: they can be mapped to a Component ontology class during the '
    'Week 2 experiment. Data/file/input/output may map to Data; shuffle and scheduling-related terms should be '
    'retained initially because they can distinguish communication and resource-management topics.'
)

doc.add_heading('6. Deliverables', level=1)
deliverables = [
    ('week1_mapreduce.py', 'Reproducible downloader and preprocessing pipeline'),
    ('issues_processed.jsonl', 'Full issue fields plus processed tokens'),
    ('issues_metadata.csv', 'Compact tabular issue metadata'),
    ('vocabulary.csv', 'Corpus and document frequencies'),
    ('document_term_matrix.npz/.mtx', 'Sparse DTM for LDA and other tools'),
    ('document_ids.json / vocabulary_terms.json', 'Matrix row and column mappings'),
    ('summary.json / download_failures.json', 'Validation summary and failure audit'),
]
for name, description in deliverables:
    p = doc.add_paragraph(style='List Bullet'); p.add_run(name).bold = True; p.add_run(f' - {description}')

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run('DSSE Assignment 3 - Week 1 - Apache MapReduce').font.size = Pt(9)

for t in doc.tables:
    header_pr = t.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement('w:tblHeader'); repeat.set(qn('w:val'), 'true'); header_pr.append(repeat)
    for row in t.rows:
        row_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement('w:cantSplit'); row_pr.append(cant_split)
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            margins = OxmlElement('w:tcMar')
            for edge in ('top', 'left', 'bottom', 'right'):
                node = OxmlElement(f'w:{edge}'); node.set(qn('w:w'), '90'); node.set(qn('w:type'), 'dxa'); margins.append(node)
            tc_pr.append(margins)

ROOT.mkdir(parents=True, exist_ok=True)
doc.save(ROOT / 'Week_1_Report_MapReduce.docx')
